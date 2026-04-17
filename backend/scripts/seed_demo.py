"""Populate the RankForge backend with a small topic-clustered corpus and run
smoke-test searches. One-off script; safe to re-run (it always adds new items,
does NOT dedupe — so only run once per fresh DB unless you want duplicates).

Usage:
    # Backend must be running (either via `npm run dev` or standalone uvicorn).
    # Override the port with SEED_PORT if needed (default 8765).
    cd z:/Rankforge-AI-Search
    .venv/Scripts/python.exe -m backend.scripts.seed_demo
"""
from __future__ import annotations

import os
import urllib.parse
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw

from backend.scripts import _http

PORT = int(os.environ.get("SEED_PORT", "8765"))
BASE = f"http://127.0.0.1:{PORT}"
ROOT = Path(__file__).resolve().parents[2]  # z:/Rankforge-AI-Search
DEMO_DIR = ROOT / "demo_data"


# ---------------------------------------------------------------------------
# Corpus generation


def _write_text(name: str, content: str) -> Path:
    path = DEMO_DIR / name
    path.write_text(content, encoding="utf-8")
    return path


def _write_docx(name: str, title: str, paragraphs: list[str]) -> Path:
    path = DEMO_DIR / name
    doc = Document()
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)
    return path


def _write_gradient_png(name: str) -> Path:
    """Make a distinctive abstract gradient PNG that CLIP can embed non-trivially."""
    path = DEMO_DIR / name
    w, h = 512, 512
    img = Image.new("RGB", (w, h))
    draw = ImageDraw.Draw(img)
    for y in range(h):
        # vertical gradient: deep blue -> magenta
        t = y / h
        r = int(30 + 200 * t)
        g = int(20 + 40 * t)
        b = int(160 - 80 * t)
        draw.line([(0, y), (w, y)], fill=(r, g, b))
    # add a warm sun-disc blob top-right to give CLIP something photographic to latch onto
    for radius in range(140, 0, -2):
        alpha = int(255 * (1 - radius / 140))
        draw.ellipse(
            [(w - 120 - radius, 80 - radius // 3),
             (w - 120 + radius, 80 + radius)],
            outline=None,
            fill=(255, 220 - radius, 120 - radius // 2),
        )
    img.save(path, "PNG")
    return path


def build_corpus() -> dict[str, Path]:
    DEMO_DIR.mkdir(parents=True, exist_ok=True)
    files: dict[str, Path] = {}

    # ML cluster
    files["neural-nets"] = _write_text(
        "neural-nets.md",
        "# Neural Networks\n\n"
        "A neural network is a function approximator made of layers of weighted sums "
        "followed by non-linearities. Training adjusts the weights via gradient descent "
        "to minimize a loss on labeled data. Deep networks stack many such layers and "
        "have shown strong performance on vision, language, and audio tasks. The "
        "universal approximation theorem proves that even one hidden layer can represent "
        "any continuous function given enough width.",
    )
    files["embeddings"] = _write_text(
        "embeddings.md",
        "# Embeddings\n\n"
        "An embedding is a dense vector representation of a discrete object. Text "
        "embeddings map sentences into Euclidean space so that semantic similarity "
        "corresponds to cosine similarity between vectors. Image embeddings do the same "
        "for pixels. CLIP jointly trains a text encoder and an image encoder so that "
        "text and image share a common embedding space, enabling cross-modal retrieval.",
    )
    files["vector-search"] = _write_text(
        "vector-search.txt",
        "Vector search finds nearest neighbors of a query vector under a distance "
        "metric, usually cosine. Exact nearest-neighbor search is O(N) per query, so "
        "production systems use approximate algorithms like HNSW or IVF-PQ. pgvector "
        "on PostgreSQL supports HNSW with a single CREATE INDEX statement and exposes "
        "an efficient <=> operator for cosine distance.",
    )
    files["transformers"] = _write_docx(
        "transformers.docx",
        "Transformer Models",
        [
            "The transformer architecture replaced recurrent networks as the default for "
            "sequence modeling tasks. Its core innovation is the self-attention layer.",
            "Self-attention computes pairwise similarity between every pair of positions "
            "in a sequence and uses those similarities to re-weight a value projection. "
            "This lets the model route information between distant tokens in constant depth.",
            "Large language models like GPT and Llama are decoder-only transformers. "
            "Embedding models like BERT and all-MiniLM are encoder-only transformers. "
            "CLIP uses a dual encoder: one transformer for text, one ViT for images.",
        ],
    )

    # Cooking cluster
    files["cooking-recipes"] = _write_text(
        "cooking-recipes.md",
        "# Dessert Recipes\n\n"
        "A classic chocolate lava cake: whisk dark chocolate and butter together, fold in "
        "egg yolks, beat the whites separately, combine, pour into ramekins, and bake at "
        "220 C for exactly seven minutes. The center should be molten when cut.\n\n"
        "Tiramisu layers espresso-soaked ladyfingers with mascarpone cream. Dusting with "
        "cocoa powder just before serving keeps the top from turning muddy.",
    )
    files["cooking-technique"] = _write_text(
        "cooking-technique.md",
        "# Maillard Reaction\n\n"
        "Browning in cooking comes from the Maillard reaction: amino acids and reducing "
        "sugars react at high temperature to form hundreds of flavor compounds. Dry the "
        "surface of a steak with paper towel before searing to get a better crust. "
        "The reaction starts around 140 C and accelerates above 150 C.",
    )

    # Finance cluster
    files["finance-investing"] = _write_text(
        "finance-investing.md",
        "# Long-term Investing\n\n"
        "Compound interest is the mechanism behind long-term wealth. Reinvested returns "
        "earn returns of their own. Over thirty years at 8% annualized, a rupee becomes "
        "ten rupees. Low-cost index funds historically beat most actively managed funds "
        "once fees are deducted.",
    )
    files["finance-indexing"] = _write_text(
        "finance-indexing.txt",
        "A Nifty index fund tracks the Nifty 50 benchmark of the largest companies "
        "listed on the National Stock Exchange of India. Expense ratios for passive "
        "index funds in India range from 0.05% to 0.20%, compared to 1.5% to 2.5% for "
        "actively managed equity mutual funds. Over a 10-year horizon, the fee "
        "differential compounds into a significant chunk of returns.",
    )

    # Image
    files["gradient-image"] = _write_gradient_png("gradient.png")

    return files


# ---------------------------------------------------------------------------
# Ingestion + linking


def ingest_all(files: dict[str, Path]) -> dict[str, str]:
    name_to_id: dict[str, str] = {}
    for name, path in files.items():
        print(f"  ingesting {name} ({path.name}) ...", end=" ", flush=True)
        resp = _http.post(
            BASE, "/content/ingest",
            body={"source_path": str(path), "title": name.replace("-", " ").title()},
        )
        name_to_id[name] = resp["id"]
        print("ok", resp["id"][:8])
    return name_to_id


def create_links(name_to_id: dict[str, str]) -> None:
    def link(src: str, dst: str) -> None:
        print(f"  {src} -> {dst}")
        _http.post(BASE, "/links", body={
            "from_id": name_to_id[src],
            "to_id": name_to_id[dst],
        })

    # ML chain: each item references embeddings, making it the hub
    link("neural-nets", "embeddings")
    link("vector-search", "embeddings")
    link("transformers", "embeddings")
    link("transformers", "neural-nets")

    # Cross-cluster pointers at embeddings as a semantic-search concept
    link("finance-indexing", "embeddings")
    link("cooking-recipes", "embeddings")


# ---------------------------------------------------------------------------
# Search tests


def run_test_queries() -> None:
    queries = [
        ("neural network", "ML cluster (semantic)"),
        ("embeddings", "hub node (semantic + FTS + pagerank)"),
        ("Nifty index fund", "exact keyword (FTS must dominate)"),
        ("dessert", "cooking cluster"),
        ("abstract gradient", "cross-modal image retrieval"),
    ]
    for q, desc in queries:
        print(f"\n[query] {q!r}  ({desc})")
        resp = _http.get(BASE, f"/content/search?q={urllib.parse.quote(q)}&limit=5")
        for i, r in enumerate(resp, 1):
            print(
                f"  #{i}  [{r['content_type']:8s}]  "
                f"final={r['final_score']:.3f}  "
                f"sem={r['semantic_similarity']:.2f}  "
                f"kw={r['keyword_match']:.2f}  "
                f"pr={r['pagerank_norm']:.2f}  "
                f"fr={r['freshness_boost']:.2f}   "
                f"{r['title']}"
            )


# ---------------------------------------------------------------------------
# Entrypoint


def main() -> None:
    _http.require_backend(BASE)
    print(f"[seed] backend OK at {BASE}")

    print("\n[seed] building corpus...")
    files = build_corpus()
    print(f"  wrote {len(files)} files to {DEMO_DIR}")

    print("\n[seed] ingesting...")
    name_to_id = ingest_all(files)

    print("\n[seed] creating links...")
    create_links(name_to_id)

    print("\n[seed] recomputing PageRank...")
    r = _http.post(BASE, "/jobs/pagerank")
    print(f"  {r}")

    print("\n[seed] running test queries...")
    run_test_queries()

    print("\n[seed] done.")


if __name__ == "__main__":
    main()
